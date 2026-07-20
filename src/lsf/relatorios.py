"""Relatório analítico do orçamento (Fase 1, item 5 do contrato).

Dois formatos, ambos gerados como string pelo motor (D6 — zero framework web):
  relatorio_csv(venda)  -> CSV ';' com decimal vírgula (conferência em planilha pt-BR)
  relatorio_html(venda) -> HTML estático

D4 na saída: itens 'estimado'/'parametrico' exibem FAIXA (±faixa_pct, default 15%)
em vez de valor seco. Rodapé traz alertas: macroetapas zeradas (aviso na Fase 1;
vira gate duro na Fase 3) e pendências de preço.
"""
from __future__ import annotations

import csv
import html
import io

from lsf.motores.orcamento import OrcamentoVenda

FAIXA_PCT_DEFAULT = 0.15


def _brl(valor: float | None) -> str:
    if valor is None:
        return ""
    return f"{valor:,.2f}".replace(",", "\0").replace(".", ",").replace("\0", ".")


def _num(valor: float | None, casas: int = 2) -> str:
    if valor is None:
        return ""
    return f"{valor:.{casas}f}".replace(".", ",")


def _tem_faixa(confianca: str | None) -> bool:
    return confianca in ("estimado", "parametrico")


def _faixa(valor: float | None, confianca: str | None, faixa_pct: float):
    """(mínimo, máximo) para valores de baixa confiança; (None, None) para 'real'."""
    if valor is None or not _tem_faixa(confianca):
        return None, None
    return valor * (1 - faixa_pct), valor * (1 + faixa_pct)


def relatorio_csv(venda: OrcamentoVenda, faixa_pct: float = FAIXA_PCT_DEFAULT) -> str:
    saida = io.StringIO()
    w = csv.writer(saida, delimiter=";", lineterminator="\n")
    w.writerow([
        "eap", "descricao", "unidade", "quantidade", "origem",
        "custo_unitario", "custo_direto", "preco_bdi",
        "preco_min", "preco_max", "confianca", "pendencia",
    ])
    for l in venda.linhas:
        pmin, pmax = _faixa(l.preco_venda, l.confianca, faixa_pct)
        w.writerow([
            l.eap_codigo, l.descricao, l.unidade, _num(l.quantidade), l.origem,
            _num(l.custo_unitario, 4), _num(l.custo_direto), _num(l.preco_venda),
            _num(pmin), _num(pmax), l.confianca or "", l.pendencia or "",
        ])
    orc = venda.orcamento
    w.writerow([])
    w.writerow(["TOTAL", f"BDI {_num(venda.bdi * 100)}%", "", "", "",
                "", _num(orc.total), _num(venda.preco_total), "", "",
                venda.confianca or "", "; ".join(orc.pendencias)])
    for codigo in orc.macroetapas_zeradas:
        w.writerow(["ALERTA", f"macroetapa {codigo} sem quantitativo (escopo vazado? R7)"])
    return saida.getvalue()


def relatorio_html(venda: OrcamentoVenda, faixa_pct: float = FAIXA_PCT_DEFAULT) -> str:
    orc = venda.orcamento
    e = html.escape

    def celula_preco(l):
        if l.preco_venda is None:
            return f'<td class="pend" colspan="1">pendente</td>'
        if _tem_faixa(l.confianca):
            pmin, pmax = _faixa(l.preco_venda, l.confianca, faixa_pct)
            return (f"<td>R$ {_brl(pmin)} – R$ {_brl(pmax)}"
                    f'<div class="ref">ref. R$ {_brl(l.preco_venda)} ±{_num(faixa_pct * 100, 0)}%</div></td>')
        return f"<td>R$ {_brl(l.preco_venda)}</td>"

    linhas_html = []
    for l in venda.linhas:
        linhas_html.append(
            f'<tr class="{e(l.confianca or "pendente")}">'
            f"<td>{e(l.eap_codigo)}</td><td>{e(l.descricao)}</td><td>{e(l.unidade)}</td>"
            f'<td class="n">{_num(l.quantidade)}</td><td>{e(l.origem)}</td>'
            f'<td class="n">{_brl(l.custo_unitario)}</td><td class="n">{_brl(l.custo_direto)}</td>'
            f"{celula_preco(l)}<td>{e(l.confianca or '—')}</td></tr>"
        )

    subtotais_html = []
    for s in orc.subtotais:
        valor = "ZERADA" if s.zerada else ("pendente" if s.custo is None else f"R$ {_brl(s.custo)}")
        subtotais_html.append(
            f"<tr><td>{e(s.eap_codigo)}</td><td>{e(s.descricao)}</td>"
            f'<td class="n">{valor}</td><td>{e(s.confianca or "—")}</td></tr>'
        )

    alertas = [
        f"Macroetapa {codigo} sem nenhum quantitativo — escopo vazado em turn-key é prejuízo (R7). "
        f"Aviso na Fase 1; bloqueia proposta a partir da Fase 3."
        for codigo in orc.macroetapas_zeradas
    ] + list(orc.pendencias)
    alertas_html = "".join(f"<li>{e(a)}</li>" for a in alertas) or "<li>nenhum</li>"

    total_html = (
        f"R$ {_brl(venda.preco_total)}" if venda.preco_total is not None
        else "NÃO FECHA — resolver pendências"
    )

    return f"""<!doctype html><html lang="pt-BR"><head><meta charset="utf-8">
<title>Orçamento analítico — {e(orc.projeto_codigo)}</title>
<style>
 body{{font-family:system-ui,sans-serif;margin:2rem;color:#222}}
 table{{border-collapse:collapse;width:100%;margin:1rem 0}}
 th,td{{border:1px solid #ccc;padding:.35rem .5rem;font-size:.9rem;text-align:left}}
 td.n{{text-align:right;font-variant-numeric:tabular-nums}}
 tr.estimado td,tr.parametrico td{{background:#fff8e6}}
 tr.pendente td,td.pend{{background:#fdecea}}
 .ref{{font-size:.75rem;color:#666}}
 .alertas{{background:#fdecea;padding:1rem;border-left:4px solid #c0392b}}
 .rodape{{font-size:.8rem;color:#666;margin-top:2rem}}
</style></head><body>
<h1>Orçamento analítico — {e(orc.projeto_codigo)}</h1>
<p>Referência {e(orc.referencia)}/{e(orc.uf or "BR")} · {"desonerado" if orc.desonerado else "não desonerado"}
 · BDI {_num(venda.bdi * 100)}% (TCU, {e(venda.parametros.confianca)}) · confiança geral: {e(venda.confianca or "—")}</p>
<table><thead><tr><th>EAP</th><th>Descrição</th><th>Un</th><th>Qtd</th><th>Origem</th>
<th>Custo unit. (R$)</th><th>Custo direto (R$)</th><th>Preço c/ BDI</th><th>Confiança</th></tr></thead>
<tbody>{"".join(linhas_html)}</tbody></table>
<h2>Subtotais por macroetapa (custo direto)</h2>
<table><thead><tr><th>EAP</th><th>Macroetapa</th><th>Custo direto</th><th>Confiança</th></tr></thead>
<tbody>{"".join(subtotais_html)}</tbody></table>
<h2>Total</h2>
<p><strong>Custo direto: {"R$ " + _brl(orc.total) if orc.total is not None else "NÃO FECHA"}
 · Preço de venda: {total_html}</strong></p>
<div class="alertas"><h2>Alertas</h2><ul>{alertas_html}</ul></div>
<p class="rodape">PRÉ-DIMENSIONAMENTO para fins de orçamento — não substitui projeto estrutural,
sondagem ou ART. Itens 'estimado'/'parametrico' exibidos em faixa ±{_num(faixa_pct * 100, 0)}% (D4).
Gerado pelo motor de orçamento LSF Veks.</p>
</body></html>"""


def exportar_mspdi(cronograma, inicio) -> str:
    """Cronograma em MSPDI (XML do MS Project) — o formato que o ProjectLibre
    importa; é a ponta automatizável da 'validação cruzada' da Fase 4 (a
    conferência visual no ProjectLibre é ação humana sobre este arquivo).

    Convenções (documentadas, não escondidas): dias CORRIDOS a partir de
    `inicio` (datetime.date); Duration em horas de jornada (dias × 8h);
    vínculos TI→Type 1 (FS), II→3 (SS), TT→0 (FF), lag em décimos de minuto de
    jornada (LagFormat 7 = dias)."""
    import datetime
    import xml.etree.ElementTree as ET

    ns = "http://schemas.microsoft.com/project"
    ET.register_namespace("", ns)
    raiz = ET.Element(f"{{{ns}}}Project")
    ET.SubElement(raiz, f"{{{ns}}}Name").text = cronograma.projeto_codigo
    ET.SubElement(raiz, f"{{{ns}}}StartDate").text = f"{inicio.isoformat()}T08:00:00"
    tarefas = ET.SubElement(raiz, f"{{{ns}}}Tasks")

    uid_por_grupo = {p.atividade.grupo: i + 1
                     for i, p in enumerate(cronograma.atividades)}
    vinculos = {p.atividade.grupo: [] for p in cronograma.atividades}
    for pred, succ, tipo, lag in (cronograma.rede or []):
        if pred in uid_por_grupo and succ in uid_por_grupo:
            vinculos[succ].append((pred, tipo, lag))

    tipo_mspdi = {"TI": "1", "II": "3", "TT": "0"}
    for p in cronograma.atividades:
        a = p.atividade
        t = ET.SubElement(tarefas, f"{{{ns}}}Task")
        ET.SubElement(t, f"{{{ns}}}UID").text = str(uid_por_grupo[a.grupo])
        ET.SubElement(t, f"{{{ns}}}ID").text = str(uid_por_grupo[a.grupo])
        ET.SubElement(t, f"{{{ns}}}Name").text = a.descricao
        ET.SubElement(t, f"{{{ns}}}Duration").text = (
            f"PT{int(a.duracao_dias * 8)}H0M0S")
        ET.SubElement(t, f"{{{ns}}}DurationFormat").text = "7"
        d0 = inicio + datetime.timedelta(days=p.es)
        d1 = inicio + datetime.timedelta(days=p.ef)
        ET.SubElement(t, f"{{{ns}}}Start").text = f"{d0.isoformat()}T08:00:00"
        ET.SubElement(t, f"{{{ns}}}Finish").text = f"{d1.isoformat()}T17:00:00"
        ET.SubElement(t, f"{{{ns}}}Critical").text = "1" if p.critica else "0"
        for pred, tipo, lag in vinculos.get(a.grupo, []):
            v = ET.SubElement(t, f"{{{ns}}}PredecessorLink")
            ET.SubElement(v, f"{{{ns}}}PredecessorUID").text = (
                str(uid_por_grupo[pred]))
            ET.SubElement(v, f"{{{ns}}}Type").text = tipo_mspdi[tipo]
            ET.SubElement(v, f"{{{ns}}}LinkLag").text = str(int(lag * 8 * 60 * 10))
            ET.SubElement(v, f"{{{ns}}}LagFormat").text = "7"

    corpo = ET.tostring(raiz, encoding="unicode")
    return '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + corpo


def romaneio_csv(romaneio) -> str:
    """Romaneio em CSV ';' (padrão pt-BR da casa), em duas seções:

    FABRICA — kit de corte por painel/perfil (o que a serra precisa);
    OBRA — sequência de montagem por nível/painel (o que o montador confere).
    """
    saida = io.StringIO()
    w = csv.writer(saida, delimiter=";")
    w.writerow(["ROMANEIO FABRICA"])
    w.writerow(["painel", "perfil", "pecas", "ml", "kg", "barras 6m",
                "perda %"])
    for item in romaneio.paineis:
        for kit in item.kits:
            w.writerow([item.painel.id, kit.perfil, kit.n_pecas,
                        str(kit.ml).replace(".", ","),
                        str(kit.kg).replace(".", ","), kit.barras,
                        str(kit.perda_pct).replace(".", ",")])
    w.writerow([])
    w.writerow(["ROMANEIO OBRA"])
    w.writerow(["nivel", "painel", "parede", "faixa (m)", "kg"])
    for item in romaneio.paineis:
        w.writerow([item.nivel_indice, item.painel.id, item.parede_id,
                    f"{item.painel.x_ini:.2f}-{item.painel.x_fim:.2f}".replace(".", ","),
                    str(item.painel.kg).replace(".", ",")])
    w.writerow([])
    w.writerow(["kg total", str(romaneio.kg_total).replace(".", ",")])
    return saida.getvalue()


def proposta_docx(venda, projeto: dict, pendencias_gate: list[str]) -> bytes:
    """Proposta comercial em .docx com identidade Veks (python-docx, MIT —
    docs/04). Documento de TRABALHO: a proposta congelada continua sendo o
    snapshot publicado em /p/<token>; este .docx serve à negociação.

    D4 na saída, como no HTML: item estimado/parametrico sai com faixa, nunca
    valor seco; gates e sondagem aparecem como seção, não rodapé morto."""
    import io as _io

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("VEKS ENGENHARIA — Proposta de Orçamento", level=0)
    doc.add_paragraph(
        f"Obra: {projeto['nome']} (cód. {projeto['codigo']}) · Cliente:"
        f" {projeto.get('cliente') or '—'}")
    doc.add_paragraph(
        f"Referência de preços: {venda.orcamento.referencia}"
        f" · UF {venda.orcamento.uf or '—'}"
        f" · {'desonerado' if venda.orcamento.desonerado else 'não desonerado'}")

    tabela = doc.add_table(rows=1, cols=3)
    cab = tabela.rows[0].cells
    cab[0].text, cab[1].text, cab[2].text = "Macroetapa", "Custo direto", "Confiança"
    for sub in venda.orcamento.subtotais:
        linha = tabela.add_row().cells
        linha[0].text = f"{sub.eap_codigo} — {sub.descricao}"
        linha[1].text = ("—" if sub.custo is None
                         else f"R$ {sub.custo:,.2f}".replace(",", "X")
                         .replace(".", ",").replace("X", "."))
        linha[2].text = sub.confianca or ("zerada" if sub.zerada else "—")

    doc.add_paragraph("")
    total = doc.add_paragraph()
    r = total.add_run(
        "PREÇO TOTAL (com BDI "
        f"{venda.bdi * 100:.2f}".replace(".", ",") + "%): "
        + ("indisponível — há pendências" if venda.preco_total is None else
           f"R$ {venda.preco_total:,.2f}".replace(",", "X")
           .replace(".", ",").replace("X", ".")))
    r.bold = True
    r.font.size = Pt(14)

    doc.add_heading("Condições e gates técnicos", level=1)
    doc.add_paragraph(
        "Este documento apresenta PRÉ-DIMENSIONAMENTO para fins de orçamento —"
        " não substitui projeto executivo, ART ou verificação estrutural.")
    if projeto.get("sondagem_pendente"):
        doc.add_paragraph(
            "• Sondagem PENDENTE: tensão de solo presumida (conservadora);"
            " confirmar por sondagem SPT antes do contrato [NBR 6122/8036].")
    for p in pendencias_gate:
        doc.add_paragraph(f"• {p}")

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
